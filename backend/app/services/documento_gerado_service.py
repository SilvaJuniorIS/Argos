import logging
from io import BytesIO
import re
import base64
from datetime import datetime

from sqlalchemy.orm import Session

from app.models.documento_gerado import DocumentoGerado
from app.models.processo_licitatorio import ProcessoLicitatorio
from app.schemas.documento_gerado import DocumentoGerarRequest
from app.services import processo_licitatorio_service
from app.services import institucional_service
from app.services.llm_gateway import (
    LLMConfigurationError,
    LLMError,
    LLMResponse,
    LLMTimeoutError,
    get_llm_gateway,
)
from app.services.prompt_engine import gerar_prompt_documento

logger = logging.getLogger("argos.documentos")


class DocumentoGeracaoError(RuntimeError):
    pass


class LLMDocumentoConfigurationError(DocumentoGeracaoError):
    pass


class LLMDocumentoTimeoutError(DocumentoGeracaoError):
    pass


def gerar_documento(db: Session, payload: DocumentoGerarRequest) -> DocumentoGerado:
    processo = processo_licitatorio_service.get_processo(db, payload.processo_id)
    if processo is None:
        raise ValueError("Processo licitatorio nao encontrado")

    tipo_original = processo.tipo_documento
    if payload.tipo_documento is not None:
        processo.tipo_documento = payload.tipo_documento

    institucional = institucional_service.get_config(db)
    prompt = gerar_prompt_documento(processo, institucional)
    logger.info(
        "Gerando documento %s para processo %s via LLM Gateway",
        processo.tipo_documento,
        processo.id,
    )

    try:
        llm_response = _gerar_texto_com_llm(prompt)
    finally:
        processo.tipo_documento = tipo_original

    documento = DocumentoGerado(
        processo_id=processo.id,
        tipo_documento=payload.tipo_documento or tipo_original,
        modelo=llm_response.model,
        conteudo=llm_response.content,
        prompt_utilizado=prompt,
    )
    db.add(documento)
    db.commit()
    db.refresh(documento)
    logger.info("Documento %s gerado e salvo para processo %s", documento.id, processo.id)
    return documento


def get_documento(db: Session, documento_id: int) -> DocumentoGerado | None:
    return db.get(DocumentoGerado, documento_id)


def atualizar_documento(db: Session, documento: DocumentoGerado, conteudo: str) -> DocumentoGerado:
    documento.conteudo = conteudo.strip()
    documento.revisado_em = datetime.now()
    db.add(documento)
    db.commit()
    db.refresh(documento)
    logger.info("Documento %s atualizado por revisao humana", documento.id)
    return documento


def exportar_docx(documento: DocumentoGerado) -> tuple[BytesIO, str]:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise DocumentoGeracaoError("Biblioteca python-docx nao instalada no backend") from exc

    doc = Document()
    _configurar_documento(doc, Inches, Pt, RGBColor)
    _adicionar_cabecalho_rodape(doc, documento, Inches, Pt)
    _adicionar_identificacao(doc, documento, WD_ALIGN_PARAGRAPH, RGBColor)

    blocos = documento.conteudo.splitlines()
    index = 0
    while index < len(blocos):
        texto = blocos[index].strip()
        if not texto:
            doc.add_paragraph("")
            index += 1
            continue
        if _parece_linha_tabela(texto):
            linhas_tabela = []
            while index < len(blocos) and _parece_linha_tabela(blocos[index].strip()):
                linhas_tabela.append(blocos[index].strip())
                index += 1
            if _adicionar_tabela_markdown(doc, linhas_tabela):
                continue
            for linha in linhas_tabela:
                _adicionar_bloco_formatado(doc, linha)
            continue
        _adicionar_bloco_formatado(doc, texto)
        index += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"argos-{documento.tipo_documento.lower()}-{documento.id}.docx"
    logger.info("Documento %s exportado em DOCX", documento.id)
    return buffer, filename


def _configurar_documento(doc, Inches, Pt, RGBColor) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in (("Title", 16), ("Heading 1", 13), ("Heading 2", 12)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(10, 35, 66)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def _adicionar_cabecalho_rodape(doc, documento: DocumentoGerado, Inches, Pt) -> None:
    section = doc.sections[0]
    institucional = getattr(documento, "_institucional", None) or _institucional_do_documento(documento)
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = True
    logo_cell = header_table.rows[0].cells[0]
    info_cell = header_table.rows[0].cells[1]

    logo_bytes = _decode_logo(getattr(institucional, "logo_base64", None))
    if logo_bytes:
        logo_run = logo_cell.paragraphs[0].add_run()
        logo_run.add_picture(BytesIO(logo_bytes), width=Inches(0.75))
    else:
        logo_cell.text = "ARGOS"

    info = info_cell.paragraphs[0]
    info.alignment = 1
    linhas = [
        getattr(institucional, "nome_orgao", None) or "Prefeitura Municipal",
        _linha_municipio(institucional),
        _linha_contato(institucional),
    ]
    for index, linha in enumerate([linha for linha in linhas if linha]):
        if index:
            info.add_run().add_break()
        run = info.add_run(linha)
        run.bold = index == 0
        run.font.size = Pt(10 if index == 0 else 8)

    footer = section.footer.paragraphs[0]
    rodape = getattr(institucional, "rodape_documentos", None) or (
        "Minuta para revisao humana especializada."
    )
    footer.text = f"{rodape} | Processo {documento.processo_id} | Documento {documento.id}"
    footer.style = doc.styles["Normal"]
    footer.runs[0].font.size = Pt(8)


def _adicionar_identificacao(doc, documento: DocumentoGerado, WD_ALIGN_PARAGRAPH, RGBColor) -> None:
    nome_documento = _nome_documento(documento.tipo_documento)
    titulo = doc.add_heading(nome_documento, level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    processo = documento.processo
    institucional = _institucional_do_documento(documento)
    secretaria = getattr(processo, "secretaria", None) or "A preencher"
    objeto = getattr(processo, "objeto", None) or "A preencher"
    municipio = _linha_municipio(institucional) or "A preencher"

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Minuta padronizada gerada com apoio do ARGOS")
    run.bold = True
    run.font.color.rgb = RGBColor(10, 35, 66)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Orgao", getattr(institucional, "nome_orgao", None) or "A preencher"),
        ("Municipio/UF", municipio),
        ("Processo", str(documento.processo_id)),
        ("Secretaria", secretaria),
        ("Objeto", objeto),
        (
            "Responsavel tecnico",
            getattr(institucional, "responsavel_tecnico", None) or "A preencher",
        ),
    ]
    for row, (label, value) in zip(table.rows, rows, strict=True):
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True

    aviso = doc.add_paragraph()
    aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aviso_run = aviso.add_run(
        "Este documento e uma minuta de apoio e deve passar por revisao tecnica, juridica "
        "e administrativa antes de qualquer uso oficial."
    )
    aviso_run.italic = True
    aviso_run.font.size = doc.styles["Normal"].font.size


def _adicionar_bloco_formatado(doc, texto: str) -> None:
    if _parece_titulo_numerado(texto):
        nivel = 2 if re.match(r"^\d+\.\d+", texto) else 1
        doc.add_heading(texto, level=nivel)
        return
    if _parece_titulo(texto):
        doc.add_heading(texto.rstrip(":"), level=1)
        return
    if texto.startswith(("- ", "• ")):
        doc.add_paragraph(texto[2:].strip(), style="List Bullet")
        return
    if re.match(r"^[a-zA-Z]\)\s+", texto):
        doc.add_paragraph(texto, style="List Bullet")
        return
    doc.add_paragraph(texto)


def _nome_documento(tipo_documento: str) -> str:
    tipo = tipo_documento.upper()
    if tipo == "ETP":
        return "Estudo Tecnico Preliminar"
    if tipo == "TR":
        return "Termo de Referencia"
    if tipo == "EDITAL":
        return "Minuta de Edital"
    return tipo_documento


def _parece_linha_tabela(texto: str) -> bool:
    return texto.startswith("|") and texto.endswith("|") and texto.count("|") >= 2


def _adicionar_tabela_markdown(doc, linhas: list[str]) -> bool:
    rows = []
    for linha in linhas:
        cells = [cell.strip() for cell in linha.strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)

    if len(rows) < 2:
        return False

    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(cols):
            value = row[col_index] if col_index < len(row) else ""
            cell = table.rows[row_index].cells[col_index]
            cell.text = value
            if row_index == 0 and cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
    return True


def _institucional_do_documento(documento: DocumentoGerado):
    session = getattr(documento, "_sa_instance_state", None)
    if session is not None:
        db = session.session
        if db is not None:
            config = institucional_service.get_config(db)
            setattr(documento, "_institucional", config)
            return config
    return None


def _linha_municipio(institucional) -> str:
    municipio = getattr(institucional, "nome_municipio", None)
    uf = getattr(institucional, "uf", None)
    if municipio and uf:
        return f"{municipio}/{uf}"
    return municipio or ""


def _linha_contato(institucional) -> str:
    partes = [
        getattr(institucional, "cnpj", None),
        getattr(institucional, "endereco", None),
        getattr(institucional, "telefone", None),
        getattr(institucional, "email", None),
        getattr(institucional, "site", None),
    ]
    return " | ".join(str(parte) for parte in partes if parte)


def _decode_logo(value: str | None) -> bytes | None:
    if not value:
        return None
    try:
        payload = value.split(",", 1)[1] if "," in value else value
        return base64.b64decode(payload)
    except Exception:
        logger.warning("Logo institucional invalido; DOCX sera gerado sem imagem")
        return None


def _parece_titulo_numerado(texto: str) -> bool:
    return bool(re.match(r"^\d+(\.\d+)*[\.\)]\s+\S+", texto)) and len(texto) <= 140


def _parece_titulo(texto: str) -> bool:
    if len(texto) > 100:
        return False
    if texto.endswith(":"):
        return True
    prefixo = texto.split(" ", 1)[0]
    return prefixo.rstrip(".").isdigit()


def _gerar_texto_com_llm(prompt: str) -> LLMResponse:
    messages = [
        {
            "role": "system",
            "content": (
                "Voce e um especialista em licitacoes publicas brasileiras e redacao tecnica "
                "de documentos baseados na Lei 14.133/2021."
            ),
        },
        {"role": "user", "content": prompt},
    ]
    try:
        return get_llm_gateway().generate_text(
            messages=messages,
            task="legal_draft",
            max_tokens=6000,
        )
    except LLMConfigurationError as exc:
        raise LLMDocumentoConfigurationError(str(exc)) from exc
    except LLMTimeoutError as exc:
        logger.warning("Timeout ao gerar documento via LLM Gateway: %s", exc)
        raise LLMDocumentoTimeoutError("Tempo excedido ao gerar documento") from exc
    except LLMError as exc:
        logger.warning("Erro ao gerar documento via LLM Gateway: %s", exc)
        raise DocumentoGeracaoError(str(exc)) from exc
    except Exception as exc:
        logger.exception("Erro inesperado ao gerar documento via LLM Gateway")
        raise DocumentoGeracaoError("Erro inesperado ao gerar documento") from exc
